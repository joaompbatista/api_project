# import libraries
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx import Document
import requests
import random
import pandas as pd
from spotipy.oauth2 import SpotifyClientCredentials
import spotipy.util as util
import os
from dotenv import load_dotenv
load_dotenv()


def get_meal(category):
    """This method requests all meals from a given category (passed as an input parameter)
        and chooses a random one from the returned dictionary"""
    meals_dict = {}
    category_url = "https://www.themealdb.com/api/json/v1/1/filter.php?c=" + category

    response = requests.get(category_url)
    if response.status_code == 200:
        meals_list = list(response.json()['meals'])
        meal_id = random.choice(meals_list)['idMeal']
    else:
        print("Sorry. Bad request.")
    meal_url = "https://www.themealdb.com/api/json/v1/1/lookup.php?i=" + meal_id
    meals = requests.get(meal_url).json()['meals'][0]
    meals_dict['Meal'] = meals['strMeal']
    meals_dict['Instructions'] = meals['strInstructions']
    meals_dict['Video'] = meals['strYoutube']
    meals_dict['Image'] = meals['strMealThumb']
    ingredients_list = []
    for i in range(1, 21):
        ingredient = meals['strIngredient' + str(i)]
        if ingredient == '':
            break
        else:
            ingredients_list.append(ingredient)
    meals_dict['Ingredients'] = ingredients_list
    return pd.DataFrame.from_dict([meals_dict])


def get_cocktail(category):
    """Document this method"""
    cocktails_dict = {}

    if category == 'Non-Alcoholic':
        # use the query to filter by non alcoholic drinks
        url = 'http://www.thecocktaildb.com/api/json/v1/1/filter.php?a=Non_Alcoholic'
    else:
        # use the query to filter by ingredient
        url = 'https://www.thecocktaildb.com/api/json/v1/1/filter.php?i=' + category
    response = requests.get(url)
    # create a list with the dictionaries for all drinks based on the category
    cocktails_list = list(response.json()['drinks'])
    # select a random cocktail
    cocktail = random.choice(cocktails_list)
    # cocktail id
    cocktail_id = cocktail['idDrink']
    # use query to lookup details about the cocktail
    url = 'http://www.thecocktaildb.com/api/json/v1/1/lookup.php?i='+cocktail_id
    response = requests.get(url)
    # cocktail name, cocktail instructions, cocktail image, cocktail video
    cocktail_name = response.json()['drinks'][0]['strDrink']  # name
    cocktail_instructions = response.json(
    )['drinks'][0]['strInstructions']  # instructions
    cocktail_image = response.json()['drinks'][0]['strDrinkThumb']  # image
    cocktail_video = response.json()['drinks'][0]['strVideo']  # video
    # cocktail ingredients
    ingredients_list = []
    for i in range(1, 16):
        ingredient = response.json(
        )['drinks'][0]['strIngredient'+str(i)]  # ingredients
        if ingredient == None:
            break
        else:
            ingredients_list.append(ingredient)

        # build dictionary with all the outputs that we want
        cocktails_dict['Cocktail'] = cocktail_name
        cocktails_dict['Instructions'] = cocktail_instructions
        cocktails_dict['Image'] = cocktail_image
        cocktails_dict['video'] = cocktail_video
        cocktails_dict['Ingredients'] = ingredients_list

    # build final dataframe
    return pd.DataFrame([cocktails_dict])


def get_playlist(category):
    """comment this method"""
    music_list = []
    username = os.environ.get("USERNAME")
    scope = os.environ.get("SCOPE")
    client_id = os.environ.get("CLIENT_ID")
    client_secret = os.environ.get("CLIENT_SECRET")
    redirect_uri = os.environ.get("REDIRECT_URI")

    token = util.prompt_for_user_token(username,
                                       scope=scope,
                                       client_id=client_id,
                                       client_secret=client_secret,
                                       redirect_uri=redirect_uri
                                       )
    endpoint_url = "https://api.spotify.com/v1/recommendations?"

    # OUR FILTERS
    limit = 5
    market = "US"
    seed_genres = category

    query = f'{endpoint_url}limit={limit}&market={market}&seed_genres={seed_genres}'

    bearer = "Bearer " + token
    response = requests.get(query,
                            headers={"Content-Type": "application/json",
                                     "Authorization": bearer})
    json_response = response.json()
    artists = []
    songs = []
    external_urls = []
    for i in json_response['tracks']:
        # print(f"\"{i['name']}\" by {i['artists'][0]['name']}")
        artists.append(i['artists'][0]['name'])
        songs.append(i['name'])
        external_urls.append(i['artists'][0]['external_urls']['spotify'])

    for music in zip(artists, songs, external_urls):
        music_list.append(music)

    # pd.DataFrame(music_list, columns=['Artist', 'Song', 'URL'])
    return random.choice(music_list)


def get_trivia(category):
    """Retrieves quotes by category of dating as a conversation starter."""
    try:
        if category == 'soft':
            url = "https://catfact.ninja/fact"
            return requests.get(url).json()['fact']
        elif category == 'romantic':
            quotes = pd.read_csv('data/love_quotes.csv')
            quotes = quotes.dropna()
            quote = quotes.sample()['Quote'].values[0]
            quote = quote.replace('“', "'").replace(
                '”', "'").replace('’', "'").replace("—", '- ')
            return quote
            #url = "https://poetrydb.org/random"
            # return "All you need is love."
        elif category == 'bold':
            url = "https://api.chucknorris.io/jokes/random"
            return requests.get(url).json()['value']
        else:
            url = "http://www.boredapi.com/api/activity/"
            return requests.get(url).json()['activity']
    except:
        print("Sorry. Bad request.")


def get_movies(category):
    movies = pd.read_csv('data/imdb_top_1000.csv')
    # Drop columns we will not use
    movies.drop(columns=['Released_Year', 'Certificate', 'Runtime', 'Overview', 'Meta_score',
                'Director', 'Star1', 'Star2', 'Star3', 'Star4', 'No_of_Votes', 'Gross'], inplace=True)
    # Filter movies dataframe to keep only the movies with "IMDB_Rating">= 8
    cond = movies['IMDB_Rating'] >= 8
    movies = movies[cond]
    # Filter movies dataframe based on category
    cond = movies['Genre'].str.contains(category)
    movies = movies[cond]
    movies.columns = ['Poster', 'Movie', 'Genre', 'IMDB_Rating']
    # Select a random movie
    return movies.sample()[['Poster', 'Movie']]


def to_doc(cocktail, meal):
    cocktail_inst = cocktail['Instructions'][0].replace(
        '\r', '').replace('\n', '')
    cocktail_ingr = ", ".join(cocktail['Ingredients'][0]).replace(
        '\r', '').replace('\n', '')
    cocktail_name = cocktail['Cocktail'][0]

    meal_inst = meal['Instructions'][0].replace('\r', '').replace('\n', '')
    meal_ingr = ", ".join(meal['Ingredients'][0]).replace(
        '\r', '').replace('\n', '')
    meal_name = meal['Meal'][0]

    # Create the document
    document = Document()

    document.add_heading("Recipes", 0)
    document.add_heading("Cocktail"+' - '+cocktail_name, level=2)

    p = document.add_paragraph()
    p = document.add_paragraph()
    run = p.add_run('Ingredients: ')
    run.bold = True
    p = document.add_paragraph(cocktail_ingr, style='List Bullet')

    p = document.add_paragraph()
    run = p.add_run('Recipe: ')
    run.bold = True
    p = document.add_paragraph(cocktail_inst, style='List Bullet')

    # Meal
    p = document.add_paragraph()
    document.add_heading("Meal"+' - '+meal_name, level=2)

    p = document.add_paragraph()
    p = document.add_paragraph()
    run = p.add_run('Ingredients: ')
    run.bold = True
    p = document.add_paragraph(meal_ingr, style='List Bullet')

    p = document.add_paragraph()
    run = p.add_run('Recipe: ')
    run.bold = True
    p = document.add_paragraph(meal_inst, style='List Bullet')

    document.save('data/recipes.docx')

    # convert("data/recipes.docx")
    # convert("data/recipes.docx", "data/recipes.pdf")
    # convert("data/")
